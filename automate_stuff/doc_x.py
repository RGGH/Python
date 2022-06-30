''' word from python '''

from docx import Document
from docx.shared import Cm

document = Document()

# Set font
style = document.styles['Normal']
style.font.name = 'Calibri'

# Add header
header_section = document.sections[0]
header = header_section.header
header_text = header.paragraphs[0]
header_text.text = "Header of document"


# Add title
document.add_heading('This is the Document Title', 0)

# Add paragraphs
p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)

document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph('first item in unordered list', style='List Bullet')

document.add_paragraph('first item in ordered list', style='List Number')

# Add image
document.add_picture('pyword.png', width=Cm(5))

records = (
    (3, '1', 'Subscribe'),
    (7, '12', 'To Python 360'),
    (4, '123', 'If you like this video!')
)

table = document.add_table(rows=1, cols=3)

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

for qty, idx, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = idx
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')

from docx2pdf import convert
docx_file = 'demo.docx'
pdf_file = 'demo.pdf'

convert(docx_file, pdf_file)


